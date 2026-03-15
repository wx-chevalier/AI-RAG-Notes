import os
import re
import shutil
import csv
import json
import uuid
from collections import OrderedDict
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

# --- Configuration ---
# 请根据实际情况修改以下路径
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOURCE_DIR = os.path.join(PROJECT_ROOT, 'sample_data')  # 原始文档目录
OUTPUT_BASE_DIR = os.path.join(PROJECT_ROOT, 'Cleaned_Knowledge_Base')
IMAGES_DIR = os.path.join(OUTPUT_BASE_DIR, 'images')
METADATA_JSON_PATH = os.path.join(OUTPUT_BASE_DIR, 'metadata.json')

# Ensure output directories exist
if os.path.exists(OUTPUT_BASE_DIR):
    shutil.rmtree(OUTPUT_BASE_DIR)
os.makedirs(OUTPUT_BASE_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)

# --- Regex Rules for Noise Filtering ---
NOISE_PATTERNS = [
    r'^产品知识库',  # 通用论坛标题
    r'.* - 汇总信息$',
    r'^共找到 \d+ 个帖子$',
    r'^帖子列表$',
    r'^详细帖子内容$',
    r'^未知标题$',
    r'^【主帖内容】$',
    r'^暂无回复$',
    r'^回复列表$',
    r'^【回复 \d+ - .*】$',
    r'^回复时间：.*$',
    r'^您的浏览器不支持 video 或 audio 标签$',
    r'^点击文件名下载附件$',
    r'^附件下载$',
    r'^.*\.zip$',
    r'^-+$',  # Separator lines like ----------------
    r'复制代码$', # Code block marker suffix
    # New patterns added
    r'^\([\d\.]+\s*(MB|KB),\s*下载次数:\s*\d+\)$', # (1.09 MB, 下载次数: 118)
    r'^\d+$', # Pure numbers like "11"
    r'^文件下载不了$'
]

# Patterns to identify image placeholders
IMAGE_PLACEHOLDER_PATTERN = re.compile(r'^(.*?\.(?:png|jpg|jpeg|gif|bmp))\s*\(.*下载次数.*\)[\s\r\n]*$')

# Pattern for Upload Timestamp
UPLOAD_TIME_PATTERN = re.compile(r'^\d{4}-\d{1,2}-\d{1,2} \d{1,2}:\d{1,2} 上传$')

def is_noise(text):
    text = text.strip()
    if not text:
        return False
    
    for pattern in NOISE_PATTERNS:
        if re.search(pattern, text):
            return True
    
    if UPLOAD_TIME_PATTERN.match(text):
        return True
        
    return False

def clean_text(text):
    if text.endswith("复制代码"):
        text = text.replace("复制代码", "")
    return text.strip()

class DocProcessor:
    def __init__(self):
        self.metadata_list = []

    def process_file(self, file_path):
        filename = os.path.basename(file_path)
        # Remove extension for display title logic, but we might want to keep robust
        # print(f"Processing: {filename}")
        
        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"Error reading {filename}: {e}")
            return

        # Prepare Clean DOCX
        clean_doc = Document()
        title = os.path.splitext(filename)[0]
        clean_doc.add_heading(title, level=1)
        
        # --- Image Extraction Phase ---
        doc_xml = doc.element.body
        image_rels_queue = []
        
        for element in doc_xml.iter():
            if element.tag.endswith('blip'):
                embed = element.get(qn('r:embed'))
                if embed:
                    image_rels_queue.append(embed)
        
        ordered_image_paths = []
        
        for rId in image_rels_queue:
            if rId in doc.part.rels:
                rel = doc.part.rels[rId]
                if "image" in rel.target_ref:
                    img_ext = rel.target_ref.split('.')[-1]
                    img_uuid = str(uuid.uuid4())[:8]
                    img_name = f"{title}_{img_uuid}.{img_ext}"
                    img_save_path = os.path.join(IMAGES_DIR, img_name)
                    
                    try:
                        with open(img_save_path, "wb") as f:
                            f.write(rel.target_part.blob)
                        ordered_image_paths.append(img_save_path)
                    except Exception as e:
                        print(f"Failed to save image {img_name}: {e}")

        ordered_image_paths.reverse() 
        
        # --- Content Processing Phase ---
        md_lines = []
        md_lines.append(f"# {title}")
        
        # Category Logic
        rel_dir = os.path.dirname(file_path)
        try:
            # 从目录路径中提取分类信息
            kb_root_name = os.path.basename(SOURCE_DIR)
            kb_root_index = rel_dir.find(kb_root_name)
            if kb_root_index != -1:
                category = rel_dir[kb_root_index + len(kb_root_name) + 1:]
                # Remove leading slash if present
                if category.startswith(os.sep):
                    category = category[1:]
            else:
                category = "Uncategorized"
        except:
            category = "Uncategorized"

        
        for para in doc.paragraphs:
            raw_text = para.text.strip()
            
            if is_noise(raw_text):
                continue
            
            img_match = IMAGE_PLACEHOLDER_PATTERN.match(raw_text)
            
            if img_match:
                if len(ordered_image_paths) > 0:
                    image_path = ordered_image_paths.pop()
                    image_filename = os.path.basename(image_path)
                    md_lines.append(f"\n![{image_filename}](images/{image_filename})\n")
                    try:
                        clean_doc.add_picture(image_path, width=Inches(6.0))
                    except FileNotFoundError:
                        clean_doc.add_paragraph(f"[Image Missing: {image_filename}]")
                else:
                    md_lines.append(f"*(Missing Image for placeholder: {img_match.group(1)})*")
                    clean_doc.add_paragraph(f"[Missing Image]")
                continue

            cleaned_t = clean_text(raw_text)
            if cleaned_t:
                clean_doc.add_paragraph(cleaned_t)
                md_lines.append(cleaned_t)
                md_lines.append("")

        while len(ordered_image_paths) > 0:
            image_path = ordered_image_paths.pop()
            image_filename = os.path.basename(image_path)
            md_lines.append(f"\n![{image_filename}](images/{image_filename}) *(Attached at end)*\n")
            clean_doc.add_paragraph("[Appendix Image]")
            clean_doc.add_picture(image_path, width=Inches(6.0))


        # --- Save Artifacts ---
        # Mirror directory structure logic could be added here, but flat list with metadata is easier for RAG.
        # Let's keep flat output for now, as metadata handles category.
        
        clean_docx_name = f"{title}_Clean.docx"
        md_name = f"{title}.md"
        
        clean_docx_path = os.path.join(OUTPUT_BASE_DIR, clean_docx_name)
        md_path = os.path.join(OUTPUT_BASE_DIR, md_name)
        
        clean_doc.save(clean_docx_path)
        
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_lines))
            
        self.metadata_list.append({
            "id": str(uuid.uuid4()),
            "filename": filename,
            "title": title,
            "category": category,
            "clean_docx_path": clean_docx_path,
            "markdown_path": md_path,
            "original_path": file_path
        })
        
        # Simple progress indicator
        if len(self.metadata_list) % 50 == 0:
            print(f"Processed {len(self.metadata_list)} files...")

    def save_metadata(self):
        with open(METADATA_JSON_PATH, 'w', encoding='utf-8') as f:
            json.dump(self.metadata_list, f, ensure_ascii=False, indent=2)
            
        print(f"Metadata saved to {METADATA_JSON_PATH}")
        print(f"Total files processed: {len(self.metadata_list)}")

def main():
    processor = DocProcessor()
    
    print(f"Starting Knowledge Base cleaning from: {SOURCE_DIR}")
    
    for root, dirs, files in os.walk(SOURCE_DIR):
        for f in files:
            if f.lower().endswith('.docx') and not f.startswith('~'):
                full_path = os.path.join(root, f)
                processor.process_file(full_path)
        
    processor.save_metadata()

if __name__ == '__main__':
    main()
